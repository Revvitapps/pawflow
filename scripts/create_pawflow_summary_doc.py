from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "/Users/mattbryan/pawflow/docs/PawFlow-Project-Summary.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(12)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("PawFlow")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(61, 58, 57)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    "A simple, smarter way for grooming and boarding businesses to run their day."
)
run.font.name = "Arial"
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(102, 96, 92)

doc.add_paragraph("")

intro = doc.add_paragraph()
intro.style = styles["Normal"]
intro.add_run(
    "PawFlow is an all-in-one system built for pet groomers, boarding businesses, "
    "and kennels. The goal is to take the daily chaos of notebooks, missed calls, "
    "text messages, sticky notes, scattered reminders, and memory, and bring it "
    "into one clean place."
)

for heading, body in [
    (
        "What It Does",
        "PawFlow helps a business manage appointments, boarding stays, customer information, "
        "pet records, grooming notes, vaccine reminders, payments, and customer messages. "
        "It also includes an AI receptionist feature that can help with missed calls, intake "
        "requests, and common customer questions.",
    ),
    (
        "Why It Matters",
        "Most pet-care businesses are still run with a mix of paper, phone calls, memory, "
        "and too many separate tools. That creates missed opportunities, staff confusion, "
        "follow-up problems, and stress for the owner. PawFlow is meant to make the business "
        "feel more organized, more professional, and easier to run every day.",
    ),
    (
        "The Bigger Opportunity",
        "This is not just about using software in one shop. The idea is to shape PawFlow "
        "inside a real grooming or kennel business first, improve it based on what really "
        "works, and then turn it into a product that can be sold to other groomers, boarding "
        "facilities, and kennels that face the same problems.",
    ),
]:
    p = doc.add_paragraph()
    r = p.add_run(heading)
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "Arial"
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    body_p = doc.add_paragraph(body)
    body_p.paragraph_format.space_after = Pt(6)

heading = doc.add_paragraph()
run = heading.add_run("A Simple Way To Explain It")
run.bold = True
run.font.size = Pt(16)
run.font.name = "Arial"
heading.paragraph_format.space_before = Pt(12)
heading.paragraph_format.space_after = Pt(4)

quote = doc.add_paragraph()
quote.paragraph_format.left_indent = Inches(0.35)
quote.paragraph_format.right_indent = Inches(0.35)
quote.paragraph_format.space_after = Pt(10)
quote_run = quote.add_run(
    '"PawFlow is a smarter, easier way to run a grooming or boarding business. '
    "It helps organize appointments, pet records, customer communication, reminders, "
    "and daily operations in one place, instead of relying on paper notes, phone calls, "
    "and memory. The goal is to first make it work really well for your business, then "
    'turn it into something we can offer to other pet-care businesses like yours."'
)
quote_run.italic = True
quote_run.font.color.rgb = RGBColor(61, 58, 57)

table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.autofit = False
table.columns[0].width = Inches(2.2)
table.columns[1].width = Inches(4.3)
hdr = table.rows[0].cells
hdr[0].text = "What She Would Be Doing"
hdr[1].text = "Why It Is Valuable"
for cell in hdr:
    set_cell_shading(cell, "DFF3F0")
    set_cell_margins(cell)
set_repeat_table_header(table.rows[0])

rows = [
    (
        "Using PawFlow in a real business",
        "It turns the prototype into something practical, tested, and useful day by day.",
    ),
    (
        "Giving feedback on what works and what does not",
        "That helps shape the system around real grooming and kennel needs instead of guesses.",
    ),
    (
        "Helping refine the customer and staff experience",
        "That makes the product stronger before offering it to other pet-care businesses.",
    ),
    (
        "Becoming part of the model we can sell later",
        "If it works well for one business first, it becomes much easier to offer to others.",
    ),
]

for left, right in rows:
    cells = table.add_row().cells
    cells[0].text = left
    cells[1].text = right
    for cell in cells:
        set_cell_margins(cell)

doc.add_paragraph("")
close = doc.add_paragraph()
close.add_run("Bottom line: ").bold = True
close.add_run(
    "PawFlow is both a tool for running a pet-care business better and a potential business "
    "model of its own. The plan is to perfect it with a real operator first, then sell that "
    "proven system to other groomers and kennels."
)

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
